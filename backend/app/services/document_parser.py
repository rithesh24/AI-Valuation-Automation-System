import logging
from pathlib import Path

import fitz  # PyMuPDF
import pytesseract
from docx import Document
from PIL import Image

from app.core.config import settings

logger = logging.getLogger(__name__)

if settings.TESSERACT_CMD:
    pytesseract.pytesseract.tesseract_cmd = settings.TESSERACT_CMD

IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png"}


class DocumentParserError(Exception):
    """Raised when a document cannot be parsed. Message is safe to show the user."""


class DocumentParser:
    def __init__(self, ocr_char_threshold: int | None = None) -> None:
        self._ocr_char_threshold = (
            ocr_char_threshold
            if ocr_char_threshold is not None
            else settings.OCR_TRIGGER_CHAR_THRESHOLD
        )

    def extract_pdf_text(self, file_path: str) -> str:
        """Extracts text from a PDF, page by page.

        A page whose extracted text falls below the OCR trigger threshold
        (D11 in docs/DECISIONS.md) is treated as a scanned image and OCR'd
        instead, so mixed text/scanned PDFs are handled transparently.
        """
        try:
            document = fitz.open(file_path)
        except Exception as exc:
            raise DocumentParserError(f"Could not open PDF '{file_path}': {exc}") from exc

        try:
            pages_text = []
            for page in document:
                text = page.get_text().strip()
                if len(text) < self._ocr_char_threshold:
                    text = self._ocr_pdf_page(page)
                pages_text.append(text)
            return "\n\n".join(pages_text)
        finally:
            document.close()

    def extract_docx_text(self, file_path: str) -> str:
        """Extracts paragraph and table text from a DOCX file, in document order."""
        try:
            document = Document(file_path)
        except Exception as exc:
            raise DocumentParserError(f"Could not open DOCX '{file_path}': {exc}") from exc

        parts = [p.text for p in document.paragraphs if p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                if any(row_text):
                    parts.append("\t".join(row_text))

        return "\n".join(parts)

    def ocr_scanned_document(self, file_path: str) -> str:
        """Runs OCR unconditionally: every page of a PDF, or a single image file
        (.jpg/.jpeg/.png — property documents may be uploaded as photos of a scan).
        """
        path = Path(file_path)
        if path.suffix.lower() in IMAGE_EXTENSIONS:
            try:
                image = Image.open(file_path)
            except Exception as exc:
                raise DocumentParserError(f"Could not open image '{file_path}': {exc}") from exc
            return self._run_ocr(image)

        try:
            document = fitz.open(file_path)
        except Exception as exc:
            raise DocumentParserError(f"Could not open PDF '{file_path}': {exc}") from exc

        try:
            return "\n\n".join(self._ocr_pdf_page(page) for page in document)
        finally:
            document.close()

    def _ocr_pdf_page(self, page: "fitz.Page") -> str:
        pixmap = page.get_pixmap(dpi=300)
        image = Image.frombytes("RGB", (pixmap.width, pixmap.height), pixmap.samples)
        return self._run_ocr(image)

    def _run_ocr(self, image: Image.Image) -> str:
        try:
            return pytesseract.image_to_string(image).strip()
        except pytesseract.TesseractNotFoundError as exc:
            logger.error("Tesseract binary not found; OCR unavailable.")
            raise DocumentParserError(
                "OCR requires Tesseract to be installed and available on PATH "
                "(or set TESSERACT_CMD in .env)."
            ) from exc
