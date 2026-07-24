import shutil
from pathlib import Path

import fitz
import pytest
from docx import Document
from PIL import Image, ImageDraw, ImageFont

from app.services.document_parser import DocumentParser, DocumentParserError

TESSERACT_AVAILABLE = shutil.which("tesseract") is not None
requires_tesseract = pytest.mark.skipif(
    not TESSERACT_AVAILABLE, reason="Tesseract binary not installed on this machine"
)


def _text_image(text: str) -> Image.Image:
    """Renders text onto a blank image using a large font, for OCR fixtures."""
    image = Image.new("RGB", (600, 150), color="white")
    draw = ImageDraw.Draw(image)
    try:
        font = ImageFont.truetype("arial.ttf", 40)
    except OSError:
        font = ImageFont.load_default()
    draw.text((10, 50), text, fill="black", font=font)
    return image


@pytest.fixture
def parser() -> DocumentParser:
    return DocumentParser()


class TestExtractPdfText:
    def test_extracts_real_text_layer(self, parser: DocumentParser, tmp_path: Path) -> None:
        pdf_path = tmp_path / "text_layer.pdf"
        document = fitz.open()
        page = document.new_page()
        page.insert_text((72, 72), "Survey Number 123 Plot Area 450 sq ft")
        document.save(pdf_path)
        document.close()

        result = parser.extract_pdf_text(str(pdf_path))

        assert "Survey Number 123" in result
        assert "450 sq ft" in result

    def test_missing_file_raises_parser_error(self, parser: DocumentParser) -> None:
        with pytest.raises(DocumentParserError):
            parser.extract_pdf_text("does_not_exist.pdf")

    @requires_tesseract
    def test_low_text_page_falls_back_to_ocr(
        self, parser: DocumentParser, tmp_path: Path
    ) -> None:
        image_path = tmp_path / "scanned_page.png"
        _text_image("PLOTNO99").save(image_path)

        pdf_path = tmp_path / "scanned.pdf"
        document = fitz.open()
        page = document.new_page()
        # No insert_text() call: the page has no text layer, only an image,
        # so page.get_text() returns well under the OCR trigger threshold.
        page.insert_image(page.rect, filename=str(image_path))
        document.save(pdf_path)
        document.close()

        result = parser.extract_pdf_text(str(pdf_path))

        assert "PLOTNO99" in result.replace(" ", "").upper()


class TestExtractDocxText:
    def test_extracts_paragraphs_and_tables(
        self, parser: DocumentParser, tmp_path: Path
    ) -> None:
        docx_path = tmp_path / "template.docx"
        document = Document()
        document.add_paragraph("Property Valuation Report")
        table = document.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "Owner Name"
        table.rows[0].cells[1].text = "Jane Doe"
        document.save(docx_path)

        result = parser.extract_docx_text(str(docx_path))

        assert "Property Valuation Report" in result
        assert "Owner Name" in result
        assert "Jane Doe" in result

    def test_missing_file_raises_parser_error(self, parser: DocumentParser) -> None:
        with pytest.raises(DocumentParserError):
            parser.extract_docx_text("does_not_exist.docx")


class TestOcrScannedDocument:
    @requires_tesseract
    def test_ocrs_direct_image_file(self, parser: DocumentParser, tmp_path: Path) -> None:
        image_path = tmp_path / "photo_of_document.png"
        _text_image("BOUNDARY").save(image_path)

        result = parser.ocr_scanned_document(str(image_path))

        assert "BOUNDARY" in result.upper()

    def test_missing_file_raises_parser_error(self, parser: DocumentParser) -> None:
        with pytest.raises(DocumentParserError):
            parser.ocr_scanned_document("does_not_exist.png")
